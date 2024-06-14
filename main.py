import os
from datetime import datetime
from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.label import Label
from kivy.uix.textinput import TextInput
from kivy.uix.button import Button
from docx import Document
from docx2pdf import convert

class InvoiceApp(App):

    def build(self):
        self.layout = BoxLayout(orientation='vertical', padding=10, spacing=10)

        self.layout.add_widget(Label(text='Rechnungsnummer:'))
        self.invoice_number_input = TextInput()
        self.layout.add_widget(self.invoice_number_input)

        self.layout.add_widget(Label(text='Rechnungsdatum (TT.MM.JJJJ):'))
        self.date_input = TextInput()
        self.layout.add_widget(self.date_input)

        self.layout.add_widget(Label(text='Leistungszeitraum:'))
        self.service_input = TextInput()
        self.layout.add_widget(self.service_input)

        self.layout.add_widget(Label(
            text='Commessa CM0184-003 Assistenza\n        di Commessa Costr 720 :'))
        self.salary1_input = TextInput()
        self.layout.add_widget(self.salary1_input)

        self.layout.add_widget(Label(
            text='Commessa CM0189-003 Assistenza\n         di Commessa Costr 718 :'))
        self.salary2_input = TextInput()
        self.layout.add_widget(self.salary2_input)

        self.layout.add_widget(Label(
            text='Commessa CM0231-003 Assistenza\n         di Commessa Costr 721 :'))
        self.salary3_input = TextInput()
        self.layout.add_widget(self.salary3_input)

        self.submit_button = Button(text='Rechnung erstellen')
        self.submit_button.bind(on_press=self.create_invoice)
        self.layout.add_widget(self.submit_button)

        return self.layout

    def replace_placeholder(self, doc, placeholder, replacement):
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, replacement)

    def create_invoice(self, instance):
        if (self.salary1_input.text == '' or self.salary2_input.text == ''
                or self.salary3_input.text == '' or self.invoice_number_input.text == ''
                or self.date_input.text == '' or self.service_input.text == ''):
            self.layout.add_widget(Label(text="Alle Werte müssen gesetzt sein"))
            return

        date = self.date_input.text
        invoice_number = self.invoice_number_input.text
        service = self.service_input.text
        salary1 = float(self.salary1_input.text)
        salary2 = float(self.salary2_input.text)
        salary3 = float(self.salary3_input.text)

        # Validate and format the date
        try:
            date = datetime.strptime(date, '%d.%m.%Y').strftime('%d.%m.%Y')
        except ValueError:
            self.layout.add_widget(Label(text='Ungültiges Datum. Bitte im Format TT.MM.JJJJ eingeben.'))
            return

        # Load the template document
        template_path = os.path.join(os.path.dirname(__file__), 'Vorlage.docx')
        doc = Document(template_path)

        # Replace placeholders with actual values
        self.replace_placeholder(doc, '{DATE}', date)
        self.replace_placeholder(doc, '{INVOICE_NUMBER}', invoice_number)
        self.replace_placeholder(doc, '{SERVICE}', service)
        self.replace_placeholder(doc, '{SALARY1}', f'{salary1:.2f}')
        self.replace_placeholder(doc, '{SALARY2}', f'{salary2:.2f}')
        self.replace_placeholder(doc, '{SALARY3}', f'{salary3:.2f}')
        brutto = salary1 + salary2 + salary3
        self.replace_placeholder(doc, '{BRUTTO}', f'{brutto:.2f}')
        self.replace_placeholder(doc, '{STEUER}', f'{(brutto * 0.19):.2f}')
        self.replace_placeholder(doc, '{NETTO}', f'{(brutto * 0.81):.2f}')

        # Save the updated document
        output_docx_path = os.path.join(os.path.dirname(__file__), f'Rechnung_Nr{invoice_number}.docx')
        doc.save(output_docx_path)

        # Convert the docx file to PDF
        output_pdf_path = os.path.join(os.path.dirname(__file__), f'Rechnung_Nr{invoice_number}.pdf')
        convert(output_docx_path, output_pdf_path)

        self.layout.add_widget(Label(text=f'Rechnung wurde erfolgreich erstellt und unter {output_docx_path} und {output_pdf_path} gespeichert!'))

if __name__ == '__main__':
    InvoiceApp().run()
